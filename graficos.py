import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import os
import zipfile
import io
from docx import Document
from docx.shared import Pt

# Cargar los datos
df = pd.read_csv('denuncias.csv')

# Cargar las stop words
stop_words_df = pd.read_csv('stop_words_es.csv')
stop_words = set(stop_words_df['stop_word'])

# Título de la aplicación
st.title("📊 Análisis y Graficos del Sistema de Gestión de Denuncias y Quejas en el Municipio de Taminango, Nariño")

# Mostrar el total de denuncias
total_denuncias = len(df)
st.header(f"📈 Total de Denuncias: {total_denuncias}")

# Panel de filtros
with st.sidebar:
    st.header("Filtros")
    
    # Filtro para municipios
    municipios = df['Municipio'].unique().tolist()
    municipios.insert(0, "Todos los datos")
    selected_municipios = st.multiselect("Seleccionar Municipios", municipios, default="Todos los datos")
    if "Todos los datos" in selected_municipios:
        selected_municipios = municipios[1:]

    # Filtro para géneros
    generos = df['Género'].unique().tolist()
    generos.insert(0, "Todos los datos")
    selected_generos = st.multiselect("Seleccionar Géneros", generos, default="Todos los datos")
    if "Todos los datos" in selected_generos:
        selected_generos = generos[1:]

    # Filtro de rango de fechas
    st.subheader("Rango de Fechas")
    start_date = st.date_input("Fecha de inicio", value=pd.to_datetime(df['Fecha'].min(), format='%Y-%m-%d'))
    end_date = st.date_input("Fecha de fin", value=pd.to_datetime(df['Fecha'].max(), format='%Y-%m-%d'))

# Aplicar filtros
df_filtered = df[(df['Municipio'].isin(selected_municipios)) & (df['Género'].isin(selected_generos))]
df_filtered = df_filtered[(df_filtered['Fecha'] >= str(start_date)) & (df_filtered['Fecha'] <= str(end_date))]

if df_filtered.empty:
    st.markdown("### 😕 No se encontraron resultados")
    st.markdown("#### No se encontraron resultados para los filtros seleccionados. Por favor, intenta con diferentes criterios de búsqueda.")
else:
    # Gráfico de distribución de denuncias por género
    st.header("👥 Distribución de Denuncias por Género")
    gender_counts = df_filtered['Género'].value_counts()
    fig, ax = plt.subplots()
    bars = gender_counts.plot(kind='bar', ax=ax, color=['skyblue', 'lightgreen', 'salmon'])
    ax.set_xlabel("Género")
    ax.set_ylabel("Cantidad de Denuncias")
    ax.set_title("Distribución de Denuncias por Género")
    for bar in bars.containers:
        ax.bar_label(bar, label_type='edge')
    st.pyplot(fig)
    st.markdown("🔍 **Explicación:** Este gráfico muestra la distribución de las denuncias según el género de la persona que realizó la denuncia. Esto ayuda a identificar si hay diferencias significativas entre las denuncias realizadas por diferentes géneros.")
    # Gráfico de distribución de denuncias por municipio
    st.header("🏙️ Distribución de Denuncias por Municipio")
    municipio_counts = df_filtered['Municipio'].value_counts()
    fig, ax = plt.subplots()
    bars = municipio_counts.plot(kind='bar', ax=ax, color='skyblue')
    ax.set_xlabel("Municipio")
    ax.set_ylabel("Cantidad de Denuncias")
    ax.set_title("Distribución de Denuncias por Municipio")
    for bar in bars.containers:
        ax.bar_label(bar, label_type='edge')
    st.pyplot(fig)
    st.markdown("🔍 **Explicación:** Este gráfico muestra la cantidad de denuncias por municipio. Es útil para identificar qué municipios tienen el mayor número de denuncias y pueden requerir más atención en términos de seguridad.")

    # Nube de palabras de los hechos denunciados
    st.header("📝 Tipos de Delitos Cometidos (Nube de Palabras)")
    text = ' '.join(df_filtered['Hecho'])
    wordcloud = WordCloud(width=800, height=400, background_color='white', stopwords=stop_words).generate(text)
    fig, ax = plt.subplots()
    ax.imshow(wordcloud, interpolation='bilinear')
    ax.axis('off')
    st.pyplot(fig)
    st.markdown("🔍 **Explicación:** La nube de palabras muestra los términos más frecuentes en las descripciones de los hechos denunciados. Ayuda a identificar los tipos de delitos más comunes.")

    # Gráfico de denuncias anónimas vs. no anónimas
    st.header("🕵️ Frecuencia de Denuncias Anónimas vs. No Anónimas")
    df_filtered['Anonima'] = df_filtered['Nombre'].apply(lambda x: 'Anónima' if x == 'Anónimo' else 'No Anónima')
    anonima_counts = df_filtered['Anonima'].value_counts()
    fig, ax = plt.subplots()
    colors = ['lightcoral', 'lightskyblue']
    anonima_counts.plot(kind='pie', ax=ax, autopct='%1.1f%%', colors=colors, labels=anonima_counts.index, startangle=90)
    ax.set_ylabel('')
    ax.set_title("Frecuencia de Denuncias Anónimas vs. No Anónimas")
    for i, (label) in enumerate(anonima_counts.index):
        ax.text(-1.5, 1-i*0.2, label, fontsize=12, color=colors[i])
    st.pyplot(fig)
    st.markdown("🔍 **Explicación:** Este gráfico de pastel muestra la proporción de denuncias anónimas frente a las no anónimas. Es útil para entender la confianza de las víctimas en el sistema de denuncias.")

    # Gráfico de distribución de denuncias por vereda
    st.header("🌾 Distribución de Denuncias por Vereda")
    vereda_counts = df_filtered['Vereda'].value_counts()
    fig, ax = plt.subplots()
    bars = vereda_counts.plot(kind='bar', ax=ax, color='lightgreen')
    ax.set_xlabel("Vereda")
    ax.set_ylabel("Cantidad de Denuncias")
    ax.set_title("Distribución de Denuncias por Vereda")
    for bar in bars.containers:
        ax.bar_label(bar, label_type='edge')
    st.pyplot(fig)
    st.markdown("🔍 **Explicación:** Este gráfico muestra la distribución de las denuncias por vereda. Ayuda a identificar si hay ciertas veredas con más denuncias y que pueden necesitar intervenciones específicas.")

    # Explicaciones adicionales y exportar datos
    st.header("📋 Resumen y Exportación de Datos")
    st.markdown("""
    - **Total de Denuncias:** El número total de denuncias presentadas en el dataset.
    - **Distribución de Denuncias por Género:** Muestra la cantidad de denuncias realizadas por cada género.
    - **Distribución de Denuncias por Municipio:** Indica en qué municipios se presentan más denuncias.
    - **Tipos de Delitos Cometidos:** Identifica los términos más frecuentes en las descripciones de los hechos denunciados.
    - **Frecuencia de Denuncias Anónimas vs. No Anónimas:** Compara la cantidad de denuncias anónimas con las no anónimas.
    - **Distribución de Denuncias por Vereda:** Muestra la cantidad de denuncias por vereda.

    Puedes exportar los datos filtrados a CSV para un análisis más detallado.
    """)
    # Botón para exportar datos
    st.download_button(
        label="📥 Descargar Datos Filtrados",
        data=df_filtered.to_csv(index=False),
        file_name='denuncias_filtradas.csv',
        mime='text/csv'
    )

    # Buscar un caso por palabra clave
    st.header("🔍 Buscar Caso por Palabra Clave")
    keyword = st.text_input("Ingrese una palabra clave para buscar casos")
    if keyword:
        matching_cases = df[df.apply(lambda row: keyword.lower() in row.astype(str).str.lower().to_string(), axis=1)]
        if not matching_cases.empty:
            st.write("### Resultados de la búsqueda")
            for index, row in matching_cases.iterrows():
                st.write(f"ID: {row['ID']}, Hecho: {row['Hecho']}")
        else:
            st.write("No se encontraron casos que coincidan con la palabra clave.")

    # Seleccionar un caso para descargar evidencias o información
    st.header("📂 Descargar Evidencias o Información de un Caso")
    selected_id = st.selectbox("Seleccionar ID de Denuncia", df_filtered["ID"].astype(str).unique())

    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("🔍 Buscar Evidencias"):
            evidencia_dir = os.path.join('evidencia', str(selected_id))
            if os.path.exists(evidencia_dir):
                st.success(f"Evidencias encontradas para la denuncia ID {selected_id}")
            else:
                st.error(f"No se encontraron evidencias para la denuncia ID {selected_id}")

    with col2:
        if st.button("📥 Descargar Evidencias 📂"):
            evidencia_dir = os.path.join('evidencia', str(selected_id))
            if os.path.exists(evidencia_dir):
                # Crear archivo ZIP en memoria
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                    for root, _, files in os.walk(evidencia_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            zip_file.write(file_path, os.path.relpath(file_path, evidencia_dir))
                
                zip_buffer.seek(0)
                st.download_button(
                    label="📥 Descargar Evidencias 📂",
                    data=zip_buffer,
                    file_name=f"{selected_id}_evidencias.zip",
                    mime="application/zip"
                )
            else:
                st.error(f"No se encontraron evidencias para la denuncia ID {selected_id}")

    with col3:
        if st.button("📄 Descargar Informe 📄"):
            matching_case = df[df["ID"] == int(selected_id)].iloc[0]
            doc = Document('FORMATO.docx')
            
            # Aquí puedes llenar tu documento con los datos del caso
            doc.add_heading(f'Informe del Caso {selected_id}', level=1)
            doc.add_paragraph(f'Nombre: {matching_case["Nombre"]}')
            doc.add_paragraph(f'Género: {matching_case["Género"]}')
            doc.add_paragraph(f'Cédula: {matching_case["Cédula"]}')
            doc.add_paragraph(f'Teléfono: {matching_case["Teléfono"]}')
            doc.add_paragraph(f'Email: {matching_case["Email"]}')
            doc.add_paragraph(f'Hecho: {matching_case["Hecho"]}')
            doc.add_paragraph(f'Departamento: {matching_case["Departamento"]}')
            doc.add_paragraph(f'Municipio: {matching_case["Municipio"]}')
            doc.add_paragraph(f'Vereda: {matching_case["Vereda"]}')
            doc.add_paragraph(f'Fecha: {matching_case["Fecha"]}')
            
            # Guardar el documento en un buffer de memoria
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📄 Descargar Informe 📄",
                data=buffer,
                file_name=f"{selected_id}_informe.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
